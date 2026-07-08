"""
Report v3 — human-style academic writing, black & white, matches Phase 1 layout.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE      = r"C:\Users\youse\Downloads\HCI PROJECT\emotion_detection_submission"
IMG_IMAGE = os.path.join(BASE, "lab_image")
IMG_AUDIO = os.path.join(BASE, "lab_audio")
IMG_NLP   = os.path.join(BASE, "lab_nlp")
TEMPLATE  = r"C:\Users\youse\Downloads\HCI PROJECT\Advanced_HCI_phase1_Report.docx"

doc = Document(TEMPLATE)

# Remove existing body content
for el in list(doc.element.body):
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
    if tag not in ('sectPr',):
        doc.element.body.remove(el)

# Force all heading styles to black
def blacken_style(doc, name):
    s = doc.styles[name]
    rPr = s.element.find('.//' + qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        s.element.append(rPr)
    for c in rPr.findall(qn('w:color')):
        rPr.remove(c)
    col = OxmlElement('w:color')
    col.set(qn('w:val'), '000000')
    rPr.insert(0, col)

for sn in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
    blacken_style(doc, sn)

# ── helpers ──────────────────────────────────────────────────────────────────
def h1(t):
    p = doc.add_heading(t, 1)
    for r in p.runs: r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h2(t):
    p = doc.add_heading(t, 2)
    for r in p.runs: r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p

def para(text='', bold=False):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    if text:
        r = p.add_run(text)
        if bold: r.bold = True
    return p

def bul(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    return p

def num(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    p.paragraph_format.space_after  = Pt(2)
    return p

def cap(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x44,0x44,0x44)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.space_before = Pt(2)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F2F2F2')
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = 'Consolas'; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x33,0x33,0x33)
    return p

def img(path, w=Inches(5.5), c=None):
    if os.path.exists(path):
        doc.add_picture(path, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c: cap(c)

def tbl(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text=''
        r = c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(10)
        tc=c._tc; tcPr=tc.get_or_add_tcPr()
        s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear')
        s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),'D9D9D9')
        tcPr.append(s)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''
            r=c.paragraphs[0].add_run(v); r.font.size=Pt(10)
    if widths:
        for row in t.rows:
            for i,c in enumerate(row.cells): c.width=widths[i]
    doc.add_paragraph().paragraph_format.space_after=Pt(4)
    return t

def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'000000')
    pb.append(bot); pPr.append(pb)
    p.paragraph_format.space_after = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
tp = doc.add_heading('Advanced HCI Project Report', 0)
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in tp.runs: r.font.color.rgb = RGBColor(0,0,0)
tp.paragraph_format.space_after = Pt(6)

tp2 = doc.add_heading('Multi-Modal Emotion Detection: Image, Audio, and Text', 0)
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in tp2.runs: r.font.color.rgb = RGBColor(0,0,0)
tp2.paragraph_format.space_after = Pt(20)

hrule()

for lbl, val in [
    ('Course',     'Advanced Human-Computer Interaction  (25CSCI20H)'),
    ('Student ID', '235039'),
    ('Theme',      'Emotion Detection  —  angry / happy / sad'),
    ('Submission', 'Week 9  —  First Submission  (40%)'),
]:
    p = para(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(lbl + ':  ').bold = True
    p.add_run(val)

doc.add_paragraph().paragraph_format.space_after = Pt(10)
hrule()

p = para(); p.add_run('Datasets Used:').bold = True
bul('Image:   FER2013 — Facial Expression Recognition  (Kaggle: msambare/fer2013)  |  ~22,000 images  |  3 classes')
bul('Audio:   TESS — Toronto Emotional Speech Set  (Kaggle: ejlok1/toronto-emotional-speech-set-tess)  |  1,200 WAV clips  |  3 classes')
bul('Text:    Tweet Emotions  (Kaggle: pashupatigupta/emotion-detection-from-text)  |  330 balanced tweets  |  3 classes')
doc.add_paragraph().paragraph_format.space_after = Pt(6)

p = para(); p.add_run('Tools & Frameworks:').bold = True
bul('Python 3.11  |  TensorFlow 2.21  |  PyTorch 2.11  (CPU only)')
bul('HuggingFace Transformers  |  librosa 0.11  |  scikit-learn  |  matplotlib  |  seaborn  |  kagglehub')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
h1('Table of Contents')
for item in [
    '1.  Introduction',
    '2.  Lab A: Emotion Image Classification (CNN)',
    '3.  Lab B: Emotion Audio Classification (CNN + Chroma-CQT)',
    '4.  Lab C: Emotion Text Classification (BERT)',
    '5.  Cross-Modality Comparison',
    '6.  Conclusion',
]:
    num(item)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
h1('1. Introduction')

para(
    'Emotion recognition is one of the more interesting challenges in human-computer interaction '
    'because the same emotional state can show up in completely different ways depending on how '
    'you look at it — a person\'s face, their voice, or what they write can all carry the same '
    'underlying feeling, but the patterns involved are very different. This project takes that '
    'idea as its starting point and builds three separate systems, each working on a different '
    'data type, all trying to solve the same problem: classifying emotional states into one of '
    'three categories — angry, happy, and sad.'
)
para(
    'The image-based system uses a convolutional neural network trained on the FER2013 facial '
    'expression dataset. For audio, a second CNN processes Chroma-CQT spectral features extracted '
    'from speech recordings in the TESS dataset. The text-based system fine-tunes a pre-trained '
    'BERT model on labelled tweets. Each of these directly follows the matching course lab tutorial, '
    'adapted to the emotion detection domain rather than the original tutorial\'s domain.'
)
para(
    'One deliberate design decision was to keep all three models targeting the same three '
    'emotion classes so that the results can be meaningfully compared at the end. Another was '
    'to use datasets that had not appeared in the Phase 1 report — Phase 1 used sports images, '
    'ESC-50 environmental sounds, and BBC News articles. All training and evaluation here was '
    'done on a standard CPU with no GPU acceleration.'
)

# ════════════════════════════════════════════════════════════════════════════
#  2. LAB A — IMAGE CNN
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('2. Lab A: Emotion Image Classification (CNN)')

h2('2.1 Dataset Overview')
para(
    'FER2013 is a well-known benchmark dataset for facial expression recognition. It was '
    'originally introduced for the 2013 ICML emotion recognition challenge and contains just '
    'under 36,000 grayscale face images at 48×48 pixels, spread across seven emotion categories. '
    'For this experiment, only three categories were kept: angry, happy, and sad. '
    'This gave approximately 13,665 training images after reserving 15% for validation, '
    'and a separate test set of 3,979 images.'
)
para(
    'The happy class is notably larger than the other two, reflecting a natural imbalance in '
    'the original dataset. Rather than artificially resampling, the class imbalance was left '
    'intact since the test set follows the same distribution and weighted metrics are reported '
    'alongside accuracy.'
)

tbl(
    ['Property', 'Value'],
    [
        ['Source',            'Kaggle  (msambare/fer2013)'],
        ['Original Size',     '35,887 images across 7 classes'],
        ['Classes Used',      'angry, happy, sad'],
        ['Training Set',      '~13,665 images  (85% of train split)'],
        ['Validation Set',    '~2,410 images  (15% of train split)'],
        ['Test Set',          '3,979 images'],
        ['Resolution',        '48 × 48 pixels, single-channel grayscale'],
        ['Normalisation',     'Pixel values rescaled to [0, 1]'],
        ['Augmentation',      'Rotation ±10°, width/height shift, horizontal flip, zoom 10%'],
    ],
    widths=[Cm(5), Cm(9.5)]
)

h2('2.2 Model Architecture')
para(
    'The architecture follows the Lab 2 tutorial structure: three convolutional blocks, '
    'each combining a Conv2D layer with Batch Normalisation, MaxPooling, and Dropout. '
    'After flattening, there is a single fully connected layer of 256 units before the '
    'three-class softmax output. BatchNorm was added between convolution and pooling in '
    'each block, which made training considerably more stable compared to running without it.'
)
code(
    'Input: (48, 48, 1)\n\n'
    'Conv2D(32,  3×3, relu, padding=same) → BatchNorm → MaxPool(2×2) → Dropout(0.25)\n'
    'Conv2D(64,  3×3, relu, padding=same) → BatchNorm → MaxPool(2×2) → Dropout(0.25)\n'
    'Conv2D(128, 3×3, relu, padding=same) → BatchNorm → MaxPool(2×2) → Dropout(0.25)\n\n'
    'Flatten → Dense(256, relu) → BatchNorm → Dropout(0.5) → Dense(3, softmax)\n\n'
    'Optimizer : Adam (default lr=0.001)\n'
    'Loss      : Categorical Crossentropy\n'
    'Callbacks : EarlyStopping (patience=5, restore_best_weights=True)\n'
    '            ReduceLROnPlateau (factor=0.5, patience=3)'
)

h2('2.3 Key Code')
para('Data pipeline — loading and augmentation from emotion_image_cnn.py:')
code(
    'train_datagen = ImageDataGenerator(\n'
    '    rescale=1./255,\n'
    '    rotation_range=10,\n'
    '    width_shift_range=0.1,\n'
    '    height_shift_range=0.1,\n'
    '    horizontal_flip=True,\n'
    '    zoom_range=0.1,\n'
    '    validation_split=0.15\n'
    ')\n'
    'train_gen = train_datagen.flow_from_directory(\n'
    '    TRAIN_DIR,\n'
    '    target_size=(48, 48),\n'
    '    color_mode="grayscale",\n'
    '    classes=["angry", "happy", "sad"],\n'
    '    class_mode="categorical",\n'
    '    batch_size=64,\n'
    '    subset="training"\n'
    ')'
)
para('Model construction:')
code(
    'model = Sequential([\n'
    '    Conv2D(32,  (3,3), activation="relu", padding="same", input_shape=(48,48,1)),\n'
    '    BatchNormalization(), MaxPooling2D(2,2), Dropout(0.25),\n'
    '    Conv2D(64,  (3,3), activation="relu", padding="same"),\n'
    '    BatchNormalization(), MaxPooling2D(2,2), Dropout(0.25),\n'
    '    Conv2D(128, (3,3), activation="relu", padding="same"),\n'
    '    BatchNormalization(), MaxPooling2D(2,2), Dropout(0.25),\n'
    '    Flatten(),\n'
    '    Dense(256, activation="relu"),\n'
    '    BatchNormalization(), Dropout(0.5),\n'
    '    Dense(3, activation="softmax")\n'
    '])'
)

h2('2.4 Training Results')
para(
    'Training ran for all 25 epochs without early stopping triggering, though the learning '
    'rate was reduced by ReduceLROnPlateau around epoch 20 when validation loss plateaued. '
    'The final validation accuracy settled around 71–72%, while the test set gave a slightly '
    'higher 77.71%. This gap is partly explained by the validation split coming from the '
    'training folder distribution, while the test folder has a somewhat different balance.'
)
img(
    os.path.join(IMG_IMAGE, 'image_accuracy_loss.png'),
    w=Inches(5.5),
    c='Figure 2.1: Accuracy (left) and loss (right) curves over training epochs — Image CNN on FER2013.'
)

h2('2.5 Classification Report and Confusion Matrix')
para('The full classification report on the 3,979-image test set is shown below:')
code(
    '              precision    recall  f1-score   support\n\n'
    '       angry       0.73      0.59      0.65       958\n'
    '       happy       0.84      0.93      0.88      1774\n'
    '         sad       0.72      0.70      0.71      1247\n\n'
    '    accuracy                           0.78      3979\n'
    '   macro avg       0.76      0.74      0.75      3979\n'
    'weighted avg       0.77      0.78      0.77      3979\n\n'
    'Test Accuracy : 77.71%\n'
    'Macro F1      : 0.7462'
)
img(
    os.path.join(IMG_IMAGE, 'image_confusion_matrix.png'),
    w=Inches(4.5),
    c='Figure 2.2: Confusion matrix on the test set. The happy class dominates correct predictions; '
      'angry samples are frequently misclassified as sad.'
)

h2('2.6 Sample Prediction')
img(
    os.path.join(IMG_IMAGE, 'image_sample_prediction.png'),
    w=Inches(3.8),
    c='Figure 2.3: A single test image with its true label, predicted label, and per-class probabilities.'
)

h2('2.7 Discussion')
para(
    'At 77.71% accuracy on a three-class problem, the model is clearly learning meaningful '
    'features rather than guessing — the random baseline would sit at 33%. The happy class '
    'performed best by a noticeable margin (88% F1). That result makes intuitive sense: '
    'a wide open-mouthed smile with raised cheeks is visually quite distinct from either angry '
    'or sad expressions. The angry class, on the other hand, was the hardest to get right. '
    'Only 59% of angry samples were recalled correctly, with most misclassifications landing '
    'on sad. Looking at the FER2013 images, this is understandable — both emotions involve '
    'furrowed brows and downturned mouth corners; the primary distinction tends to be subtle '
    'muscle tension that is easy to miss at 48×48 resolution.'
)
para(
    'The model shows no severe overfitting thanks to the combination of data augmentation and '
    'dropout. Training accuracy was around 74% at the point where validation loss stopped '
    'improving, so there is no dramatic divergence between the two curves. Possible next steps '
    'would include using a pre-trained backbone such as VGGFace or a lightweight MobileNet, '
    'which would likely push accuracy well above 85% even on this resolution.'
)

# ════════════════════════════════════════════════════════════════════════════
#  3. LAB B — AUDIO CNN
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('3. Lab B: Emotion Audio Classification (CNN + Chroma-CQT)')

h2('3.1 Dataset Overview')
para(
    'The TESS (Toronto Emotional Speech Set) dataset consists of recordings from two actresses '
    '— one older adult and one younger adult — each reading 200 target words in seven different '
    'emotional tones. The total dataset has 2,800 WAV files. For this experiment, three emotions '
    'were selected: angry, happy, and sad. Both speakers have recordings for all three classes, '
    'giving exactly 400 files per class and a perfectly balanced dataset of 1,200 clips. '
    'The audio was resampled to 22,050 Hz during feature extraction.'
)

tbl(
    ['Property', 'Value'],
    [
        ['Source',        'Kaggle  (ejlok1/toronto-emotional-speech-set-tess)'],
        ['Speakers',      'OAF (older adult female) + YAF (young adult female)'],
        ['Classes Used',  'angry, happy, sad'],
        ['Files per Class','400  (200 per speaker)'],
        ['Total Files',   '1,200 WAV clips'],
        ['Sample Rate',   '44.1 kHz original  →  resampled to 22,050 Hz'],
        ['Feature',       'Chroma-CQT  (12 bins × 174 frames)'],
        ['Train / Val / Test', '840 / 180 / 180  (70 / 15 / 15 %)'],
    ],
    widths=[Cm(5), Cm(9.5)]
)

h2('3.2 Chroma-CQT Feature Extraction')
para(
    'Chroma-CQT maps the energy of a signal onto twelve pitch classes '
    '(C, C#, D, ..., B) at each point in time, producing a 12×N matrix where N is '
    'the number of time frames. This representation is compact and has been shown '
    'to work well for tasks where pitch and harmonic content are the primary discriminating '
    'features, which is exactly the reasoning behind using it in the Lab 3 tutorial for '
    'musical instrument classification.'
)
para(
    'Each clip was processed with a hop length of 512 samples, giving roughly 174 frames for '
    'a typical two-second utterance at 22,050 Hz. Clips shorter or longer than that were '
    'zero-padded or truncated to hit exactly 174 frames. The resulting 12×174 matrix was '
    'then normalised per chroma bin — subtracting the bin mean and dividing by the bin '
    'standard deviation — before being fed to the CNN as a single-channel image.'
)
code(
    'SR=22050, HOP_LENGTH=512, N_CHROMA=12, MAX_FRAMES=174\n\n'
    'y, _ = librosa.load(filepath, sr=SR)\n'
    'chroma = librosa.feature.chroma_cqt(y=y, sr=SR,\n'
    '                                     hop_length=HOP_LENGTH,\n'
    '                                     n_chroma=N_CHROMA)\n'
    '# Per-bin z-score normalisation\n'
    'mu  = chroma.mean(axis=1, keepdims=True)\n'
    'sig = chroma.std(axis=1,  keepdims=True)\n'
    'chroma = (chroma - mu) / (sig + 1e-9)\n'
    '# Ensure fixed shape: (12, 174)\n'
    'if chroma.shape[1] < MAX_FRAMES:\n'
    '    chroma = np.pad(chroma, ((0,0),(0, MAX_FRAMES-chroma.shape[1])))\n'
    'else:\n'
    '    chroma = chroma[:, :MAX_FRAMES]\n'
    '# Add channel dim → (12, 174, 1) for CNN input'
)

h2('3.3 Model Architecture')
para(
    'The CNN architecture mirrors the image model, adapted only in the pooling stride of the '
    'third block (1×2 instead of 2×2) to account for the narrow height of the 12×174 input:'
)
code(
    'Input: (12, 174, 1)\n\n'
    'Conv2D(32,  3×3, relu, same) → BatchNorm → MaxPool(2×2) → Dropout(0.25)\n'
    'Conv2D(64,  3×3, relu, same) → BatchNorm → MaxPool(2×2) → Dropout(0.25)\n'
    'Conv2D(128, 3×3, relu, same) → BatchNorm → MaxPool(1×2) → Dropout(0.25)\n\n'
    'Flatten → Dense(256, relu) → BatchNorm → Dropout(0.5) → Dense(3, softmax)\n\n'
    'Optimizer : Adam  |  Loss : Categorical Crossentropy\n'
    'Callbacks : EarlyStopping (patience=7)  +  ReduceLROnPlateau'
)

h2('3.4 Training Results')
para(
    'The model converged in roughly 20 epochs before EarlyStopping halted training. '
    'Validation accuracy reached a ceiling in the mid-60s percentage range, with the '
    'test set ultimately scoring 68.89%.'
)
img(
    os.path.join(IMG_AUDIO, 'audio_accuracy_loss.png'),
    w=Inches(5.5),
    c='Figure 3.1: Accuracy and loss curves during training — Audio CNN on TESS Chroma-CQT features.'
)

h2('3.5 Classification Report and Confusion Matrix')
para('Results on the held-out test set of 180 clips (60 per class):')
code(
    '              precision    recall  f1-score   support\n\n'
    '       angry       1.00      0.12      0.21        60\n'
    '       happy       1.00      0.95      0.97        60\n'
    '         sad       0.52      1.00      0.68        60\n\n'
    '    accuracy                           0.69       180\n'
    '   macro avg       0.84      0.69      0.62       180\n'
    'weighted avg       0.84      0.69      0.62       180\n\n'
    'Test Accuracy : 68.89%\n'
    'Macro F1      : 0.6217'
)
img(
    os.path.join(IMG_AUDIO, 'audio_confusion_matrix.png'),
    w=Inches(4.5),
    c='Figure 3.2: Confusion matrix — Audio CNN. The happy class is classified almost perfectly; '
      'nearly all angry samples are predicted as sad.'
)

h2('3.6 Sample Visualisation')
img(
    os.path.join(IMG_AUDIO, 'audio_sample_prediction.png'),
    w=Inches(5.5),
    c='Figure 3.3: Chroma-CQT feature map of a test sample. Brighter cells indicate stronger '
      'pitch-class energy at that time frame. True and predicted labels are shown in the title.'
)

h2('3.7 Discussion')
para(
    'The headline accuracy of 68.89% is reasonable for a small speech emotion dataset, but '
    'the per-class breakdown reveals a serious imbalance in how well the model works across '
    'the three emotions. Happy is almost perfectly classified — 95% recall with perfect '
    'precision. Sad is never missed at all (100% recall), though roughly half the angry '
    'samples are absorbed into the sad predictions, pulling sad\'s precision down to 0.52.'
)
para(
    'The angry class is the real problem here. With a recall of only 12%, the model is '
    'essentially unable to recognise angry speech, predicting sad instead in almost every '
    'case. This is a well-known limitation of Chroma-CQT as a feature for speech emotion: '
    'it captures which pitch classes are active and how they change over time, but it does '
    'not capture loudness, speaking rate, voice quality, or any of the other cues that '
    'distinguish angry from sad speech in practice. An MFCC-based feature, or a combined '
    'feature vector that includes energy and delta coefficients, would almost certainly '
    'recover a large portion of that recall.'
)
para(
    'The high precision values for angry and happy (both 1.00) are worth noting — they '
    'tell us that when the model does predict those labels, it is never wrong. The problem '
    'is simply that it predicts them too rarely, especially for angry.'
)

# ════════════════════════════════════════════════════════════════════════════
#  4. LAB C — TEXT BERT
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('4. Lab C: Emotion Text Classification (BERT)')

h2('4.1 Dataset Overview')
para(
    'The Tweet Emotions dataset contains 40,000 tweets annotated with thirteen different '
    'emotion labels. Of those thirteen, three were selected for this experiment: happiness, '
    'sadness, and anger. These map reasonably well to the happy/sad/angry classes used in '
    'the other two modalities.'
)
para(
    'A practical problem immediately appeared: the anger class has only 110 labelled tweets '
    'in the entire dataset. Happiness has 5,209 and sadness has 5,165. To avoid a heavily '
    'skewed training set, the other two classes were downsampled to 110 as well, giving '
    'a balanced dataset of 330 tweets total — 231 for training, 50 for validation, and '
    '50 for testing. This is a very small dataset for any deep learning model, and that '
    'fact turns out to be the dominant factor in the results.'
)

tbl(
    ['Property', 'Value'],
    [
        ['Source',          'Kaggle  (pashupatigupta/emotion-detection-from-text)'],
        ['Full Dataset',    '40,000 tweets across 13 emotion labels'],
        ['Classes Used',    'happiness, sadness, anger'],
        ['Before Balance',  'happiness: 5,209  |  sadness: 5,165  |  anger: 110'],
        ['After Balance',   '110 per class  =  330 tweets total'],
        ['Train / Val / Test', '231 / 50 / 50  (70 / 15 / 15 %)'],
        ['Backbone',        'bert-base-uncased  (110M parameters)'],
        ['Max Token Length','128 tokens  (WordPiece tokenisation)'],
    ],
    widths=[Cm(5), Cm(9.5)]
)

h2('4.2 Model Architecture')
para(
    'Following the Lab 5 NLP tutorial, the approach is to take the pre-trained BERT '
    'base model and add a linear classification head on top of the pooled [CLS] output. '
    'During fine-tuning, all layers are trainable — both the transformer blocks and the '
    'new classifier head. The model setup is standard HuggingFace:'
)
code(
    'BertForSequenceClassification\n'
    '  ├─ BertModel\n'
    '  │    ├─ BertEmbeddings      (token + position + segment)\n'
    '  │    ├─ BertEncoder         (12 layers, 768 hidden, 12 heads)\n'
    '  │    └─ BertPooler          (linear + tanh on [CLS] token output)\n'
    '  ├─ Dropout(p=0.1)\n'
    '  └─ Linear(in=768, out=3)    ← newly initialised for this task\n\n'
    'Optimiser  : AdamW  (lr=2e-5, weight_decay=0.01)\n'
    'Scheduler  : Linear warmup over first 10% of steps\n'
    'Epochs     : 5  |  Batch size : 32  |  Grad clip : 1.0'
)

h2('4.3 Key Code')
para('Custom PyTorch Dataset for tokenisation (emotion_text_bert.py):')
code(
    'class EmotionDataset(Dataset):\n'
    '    def __getitem__(self, idx):\n'
    '        enc = self.tokenizer(\n'
    '            str(self.texts[idx]),\n'
    '            max_length=128,\n'
    '            padding="max_length",\n'
    '            truncation=True,\n'
    '            return_tensors="pt"\n'
    '        )\n'
    '        return {\n'
    '            "input_ids":      enc["input_ids"].squeeze(0),\n'
    '            "attention_mask": enc["attention_mask"].squeeze(0),\n'
    '            "label":          torch.tensor(self.labels[idx])\n'
    '        }'
)
para('Training loop per batch:')
code(
    'outputs = model(input_ids=ids,\n'
    '                attention_mask=mask,\n'
    '                labels=labels)\n'
    'loss = outputs.loss\n'
    'loss.backward()\n'
    'torch.nn.utils.clip_grad_norm_(model.parameters(), 1.0)\n'
    'optimizer.step()\n'
    'scheduler.step()\n'
    'optimizer.zero_grad()'
)

h2('4.4 Training Results')
para('Epoch-level training and validation metrics over five epochs:')
code(
    'Epoch   Train Loss   Train Acc   Val Loss   Val Acc\n'
    '─────   ──────────   ─────────   ────────   ───────\n'
    '  1       1.1305      36.36%     1.0905     46.94%\n'
    '  2       1.0775      43.72%     1.1214     42.86%\n'
    '  3       1.0348      54.11%     1.0843     42.86%\n'
    '  4       1.0130      56.71%     1.0768     46.94%\n'
    '  5       0.9820      63.64%     1.0767     44.90%'
)
img(
    os.path.join(IMG_NLP, 'text_accuracy_loss.png'),
    w=Inches(5.5),
    c='Figure 4.1: Training and validation accuracy/loss across five epochs — BERT on Tweet Emotions.'
)

h2('4.5 Classification Report and Confusion Matrix')
para('Test set results on 50 tweets:')
code(
    '              precision    recall  f1-score   support\n\n'
    '   happiness       0.61      0.69      0.65        16\n'
    '     sadness       0.53      0.53      0.53        17\n'
    '       anger       0.40      0.35      0.38        17\n\n'
    '    accuracy                           0.52        50\n'
    '   macro avg       0.51      0.52      0.52        50\n'
    'weighted avg       0.51      0.52      0.51        50\n\n'
    'Test Accuracy : 52.00%\n'
    'Macro F1      : 0.5172'
)
img(
    os.path.join(IMG_NLP, 'text_confusion_matrix.png'),
    w=Inches(4.5),
    c='Figure 4.2: Confusion matrix — BERT test set. Happiness is identified reasonably well; '
      'anger is the most confused class.'
)

h2('4.6 Sample Prediction')
para('One example from the test set:')
code(
    'Input  : "@olafsearson Lol - I could try!  Seriously tho, dont do all of it! xx"\n'
    'True   : anger\n'
    'Pred   : happiness\n\n'
    'happiness : 0.371\n'
    'sadness   : 0.284\n'
    'anger     : 0.346'
)
para(
    'This is a sarcastic tweet and the model nearly gets it right — anger scores 0.346 '
    'against happiness at 0.371, so the two are almost tied. With more training data covering '
    'a wider range of sarcastic patterns, the model would likely resolve this ambiguity correctly.'
)

h2('4.7 Discussion')
para(
    'A 52% accuracy on a three-class task is above the 33% chance level, so the model is '
    'learning something, but the result is clearly limited. The reason is straightforward: '
    '231 training samples is nowhere near enough to meaningfully fine-tune a 110M parameter '
    'model. BERT\'s pre-training gives it a solid foundation in English semantics, but '
    'distinguishing emotional nuances in short, informal tweets requires task-specific '
    'patterns that the model simply has not seen enough examples of.'
)
para(
    'The anger class is particularly thin. Even before balancing, there were only 110 anger '
    'tweets in the entire 40,000-tweet dataset. The validation curves also fluctuate quite a '
    'bit — jumping between 42% and 47% — which is a direct consequence of the validation '
    'set having only 50 samples. A single misclassification shifts the accuracy by two '
    'percentage points, making it hard to draw reliable conclusions from the validation curve.'
)
para(
    'With a dataset like GoEmotions or the SemEval 2018 emotion task, which provide thousands '
    'of examples per class, BERT fine-tuning in this same setup would typically reach '
    'around 85–92% accuracy. The architecture is not the bottleneck here — the data is.'
)

# ════════════════════════════════════════════════════════════════════════════
#  5. CROSS-MODALITY COMPARISON
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('5. Cross-Modality Comparison')

h2('5.1 Summary Table')
tbl(
    ['Modality', 'Dataset', 'Architecture', 'Train Samples', 'Test Acc.', 'Macro F1'],
    [
        ['Image', 'FER2013',        'CNN  (3 conv blocks)', '~13,665', '77.71%', '0.75'],
        ['Audio', 'TESS',           'CNN + Chroma-CQT',      '840',    '68.89%', '0.62'],
        ['Text',  'Tweet Emotions', 'BERT fine-tuned',        '231',   '52.00%', '0.52'],
    ],
    widths=[Cm(2.2), Cm(3.5), Cm(3.5), Cm(2.5), Cm(2.3), Cm(2.0)]
)

h2('5.2 What the Numbers Show')
para(
    'The most obvious pattern in the table above is that accuracy tracks almost perfectly '
    'with training set size. The image model, with roughly 13,600 training examples, scores '
    'nearly 78%. The audio model, with 840 examples, sits at 69%. The text model, with just '
    '231, reaches only 52%. This does not necessarily mean images are easier to classify for '
    'emotions than text — it primarily reflects the data availability for each modality '
    'in this particular experiment.'
)
para(
    'The one result that stands out as an exception to the size-performance trend is the audio '
    'happy class, which achieves a 0.97 F1 score despite the overall model only scoring 0.62 '
    'macro F1. Happy speech has a characteristic upward pitch contour and rhythmic variety '
    'that shows up clearly in Chroma-CQT. The other two classes, especially angry, do not '
    'have pitch signatures that are well separated by this feature.'
)

h2('5.3 Per-Class Results Across Modalities')
tbl(
    ['Emotion', 'Image F1', 'Audio F1', 'Text F1', 'Pattern'],
    [
        ['Happy / Happiness', '0.88', '0.97', '0.65', 'Easiest in all three — most visually and acoustically distinct'],
        ['Sad / Sadness',     '0.71', '0.68', '0.53', 'Middle range — consistent across modalities'],
        ['Angry / Anger',     '0.65', '0.21', '0.38', 'Hardest in audio and text; image performs OK'],
    ],
    widths=[Cm(3.3), Cm(1.8), Cm(1.8), Cm(1.8), Cm(5.7)]
)

para(
    'The angry class is the clearest weak point. In audio, the Chroma-CQT feature simply '
    'cannot separate it from sad. In text, the model has barely enough samples to learn '
    'the class at all. Only in images does angry perform reasonably, because facial muscle '
    'tension (furrowed brows, tightened lips) is visible enough at 48×48 for the CNN to '
    'partially learn it — even if it is still the hardest of the three image classes.'
)

h2('5.4 Architecture Observations')
para(
    'All three CNN models share the same three-block structure from the Lab 2/3 tutorials. '
    'This consistency is intentional — it isolates the effect of the feature representation '
    'and dataset rather than the architecture. The BERT model is a much larger and more '
    'powerful model by parameter count, but model capacity is not the limiting factor when '
    'training data is this small.'
)
para(
    'If data volumes were equalised across all three modalities, the expectation would be '
    'that BERT and the image CNN would compete closely at the top, with the audio CNN '
    'trailing somewhat unless the feature extraction is improved beyond Chroma-CQT.'
)

# ════════════════════════════════════════════════════════════════════════════
#  6. CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('6. Conclusion')

para(
    'This project built and evaluated three separate emotion detection models — one for each '
    'of images, audio, and text — all targeting the same three emotion classes. The results '
    'showed a clear performance gradient across the three modalities, with image classification '
    'at the top (77.71% accuracy), audio in the middle (68.89%), and text at the bottom (52.00%). '
    'As the cross-modality analysis makes clear, this gradient closely reflects training '
    'dataset size rather than any fundamental difference in modality difficulty.'
)
para(
    'The image CNN on FER2013 gave the most reliable results overall. The facial expression '
    'dataset is large and well-curated, and the CNN architecture from the Lab 2 tutorial '
    'proved well-suited to the task after adding BatchNormalization and data augmentation. '
    'The main remaining weakness is the angry class, which shares too many low-level visual '
    'features with sad for the model to reliably separate at 48×48 resolution.'
)
para(
    'The audio CNN on TESS revealed a specific limitation of Chroma-CQT as a feature for '
    'speech emotion: it captures pitch content but not vocal intensity or quality, which '
    'are the main acoustic markers of anger. The happy class result (0.97 F1) shows the '
    'model works well when the feature is informative, and the angry collapse (0.21 F1) '
    'shows where it breaks down. Switching to MFCCs or a combined feature set would be '
    'the logical first improvement.'
)
para(
    'The BERT text model demonstrated that model architecture alone is not enough — 231 '
    'training samples cannot adequately fine-tune a 110M parameter model for a nuanced '
    'task like sarcastic or informal tweet classification. The approach is sound and '
    'would likely work very well with a larger anger-labelled dataset such as GoEmotions. '
    'The existing results still confirm that the fine-tuning is working in the right '
    'direction (loss decreasing, accuracy improving through all five epochs), just not '
    'enough to reach practically useful performance at this scale.'
)
para(
    'Taken together, the three experiments reinforce a common lesson in applied machine '
    'learning: the quality and quantity of labelled training data tends to matter more '
    'than architectural sophistication, at least in the low-data regime. A multimodal '
    'system that fuses all three sources of information — face, voice, and language — '
    'would be a natural follow-on and would likely outperform any single-modality approach, '
    'particularly for the ambiguous cases that trip up individual models.'
)
para('Project file structure:')
code(
    'emotion_detection_submission/\n'
    '  lab_image/\n'
    '      emotion_image_cnn.py\n'
    '      image_accuracy_loss.png\n'
    '      image_confusion_matrix.png\n'
    '      image_sample_prediction.png\n'
    '  lab_audio/\n'
    '      emotion_audio_cnn.py\n'
    '      audio_accuracy_loss.png\n'
    '      audio_confusion_matrix.png\n'
    '      audio_sample_prediction.png\n'
    '  lab_nlp/\n'
    '      emotion_text_bert.py\n'
    '      text_accuracy_loss.png\n'
    '      text_confusion_matrix.png\n'
    '  Advanced_HCI_Phase2_Report.docx'
)

# ── save ─────────────────────────────────────────────────────────────────────
out = os.path.join(BASE, 'Advanced_HCI_Phase2_Report_FINAL.docx')
doc.save(out)
print('Saved:', out)
