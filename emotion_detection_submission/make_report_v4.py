"""Report v4 — human-rewritten prose, B&W, Phase 1 template layout."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE      = r"C:\Users\youse\Downloads\HCI PROJECT\emotion_detection_submission"
IMG_IMAGE = os.path.join(BASE, "lab_image")
IMG_AUDIO = os.path.join(BASE, "lab_audio")
IMG_NLP   = os.path.join(BASE, "lab_nlp")
TEMPLATE  = r"C:\Users\youse\Downloads\HCI PROJECT\Advanced_HCI_phase1_Report.docx"

doc = Document(TEMPLATE)
for el in list(doc.element.body):
    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
    if tag != 'sectPr':
        doc.element.body.remove(el)

def blacken(doc, name):
    s = doc.styles[name]
    rPr = s.element.find('.//' + qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr'); s.element.append(rPr)
    for c in rPr.findall(qn('w:color')): rPr.remove(c)
    col = OxmlElement('w:color'); col.set(qn('w:val'), '000000'); rPr.insert(0, col)

for sn in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']: blacken(doc, sn)

def h1(t):
    p = doc.add_heading(t, 1)
    for r in p.runs: r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    return p

def h2(t):
    p = doc.add_heading(t, 2)
    for r in p.runs: r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    return p

def para(text=''):
    p = doc.add_paragraph(); p.style = doc.styles['Normal']
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.space_before = Pt(0)
    if text: p.add_run(text)
    return p

def bul(text):
    p = doc.add_paragraph(style='List Bullet'); p.add_run(text)
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(0)
    return p

def num(text):
    p = doc.add_paragraph(style='List Number'); p.add_run(text)
    p.paragraph_format.space_after = Pt(2); return p

def cap(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x44,0x44,0x44)
    p.paragraph_format.space_after = Pt(10); p.paragraph_format.space_before = Pt(2)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F2F2F2')
    pPr.append(shd)
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x33,0x33,0x33)
    return p

def img(path, w=Inches(5.5), c=None):
    if os.path.exists(path):
        doc.add_picture(path, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c: cap(c)

def tbl(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text=''
        r = c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(10)
        tc=c._tc; tcPr=tc.get_or_add_tcPr()
        s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear')
        s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),'D9D9D9'); tcPr.append(s)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''
            r=c.paragraphs[0].add_run(v); r.font.size=Pt(10)
    if widths:
        for row in t.rows:
            for i,c in enumerate(row.cells): c.width=widths[i]
    doc.add_paragraph().paragraph_format.space_after=Pt(4)
    return t

def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr(); pb = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom'); bot.set(qn('w:val'),'single')
    bot.set(qn('w:sz'),'6'); bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'000000')
    pb.append(bot); pPr.append(pb); p.paragraph_format.space_after = Pt(10)

# ── TITLE PAGE ───────────────────────────────────────────────────────────────
tp = doc.add_heading('Advanced HCI Project Report', 0)
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in tp.runs: r.font.color.rgb = RGBColor(0,0,0)
tp.paragraph_format.space_after = Pt(6)
tp2 = doc.add_heading('Multi-Modal Emotion Detection: Image, Audio, and Text', 0)
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in tp2.runs: r.font.color.rgb = RGBColor(0,0,0)
tp2.paragraph_format.space_after = Pt(20)
hrule()
for lbl,val in [
    ('Course','Advanced Human-Computer Interaction  (25CSCI20H)'),
    ('Student ID','235039'),
    ('Theme','Emotion Detection  —  angry / happy / sad'),
    ('Submission','Week 9  —  First Submission  (40%)'),
]:
    p=para(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(lbl+':  ').bold=True; p.add_run(val)
doc.add_paragraph().paragraph_format.space_after=Pt(10)
hrule()
p=para(); p.add_run('Datasets Used:').bold=True
bul('Image:   FER2013 — Facial Expression Recognition  (Kaggle: msambare/fer2013)  |  ~22,000 images  |  3 classes')
bul('Audio:   TESS — Toronto Emotional Speech Set  (Kaggle: ejlok1/toronto-emotional-speech-set-tess)  |  1,200 WAV clips  |  3 classes')
bul('Text:    Tweet Emotions  (Kaggle: pashupatigupta/emotion-detection-from-text)  |  330 balanced tweets  |  3 classes')
doc.add_paragraph().paragraph_format.space_after=Pt(6)
p=para(); p.add_run('Tools & Frameworks:').bold=True
bul('Python 3.11  |  TensorFlow 2.21  |  PyTorch 2.11  (CPU only)')
bul('HuggingFace Transformers  |  librosa 0.11  |  scikit-learn  |  matplotlib  |  seaborn  |  kagglehub')
doc.add_page_break()

# ── TOC ──────────────────────────────────────────────────────────────────────
h1('Table of Contents')
for item in ['1.  Introduction','2.  Lab A: Emotion Image Classification (CNN)',
             '3.  Lab B: Emotion Audio Classification (CNN + Chroma-CQT)',
             '4.  Lab C: Emotion Text Classification (BERT)',
             '5.  Cross-Modality Comparison','6.  Conclusion']:
    num(item)
doc.add_page_break()

# ════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ════════════════════════════════════════════════════════════
h1('1. Introduction')

para("Emotion recognition is genuinely one of the harder problems in human-computer interaction, and what makes it interesting is that the same feeling can show up in completely different ways depending on which signal you're looking at. A person's face, their voice, and the words they type can all carry anger or happiness — but the underlying patterns are totally different from each other. This project starts from that observation and builds three separate systems, one per data type, each trying to classify emotional states into one of three categories: angry, happy, and sad.")

para("The image-based system is a convolutional neural network trained on the FER2013 facial expression dataset. For audio, a second CNN takes Chroma-CQT spectral features pulled from speech recordings in the TESS dataset. Text goes through a fine-tuned BERT model trained on labelled tweets. Each of these follows the corresponding course lab tutorial pretty closely, just adapted to the emotion detection domain instead of whatever the original tutorial was doing.")

para("One deliberate choice was keeping all three models pointed at the same three emotion classes — that way you can actually compare the results at the end rather than talking past each other. Also worth mentioning: the datasets here are all different from Phase 1, which used sports images, ESC-50 environmental sounds, and BBC News articles. Everything was trained and evaluated on a standard CPU, no GPU.")

# ════════════════════════════════════════════════════════════
#  2. LAB A — IMAGE
# ════════════════════════════════════════════════════════════
doc.add_page_break()
h1('2. Lab A: Emotion Image Classification (CNN)')

h2('2.1 Dataset Overview')
para("FER2013 is probably the most commonly cited benchmark for facial expression recognition. It originally came out of the 2013 ICML emotion recognition challenge and has just under 36,000 grayscale face images, all at 48×48 pixels, spread across seven emotion categories. For this experiment only three of those categories were kept — angry, happy, and sad — which left roughly 13,665 training images after setting aside 15% for validation, plus a separate test set of 3,979 images.")
para("The happy class is noticeably larger than the other two. That's just a natural quirk of the original dataset rather than anything introduced here. I decided to leave the imbalance alone rather than resampling artificially, since the test set follows the same distribution anyway and weighted metrics get reported alongside plain accuracy.")

tbl(['Property','Value'],[
    ['Source','Kaggle  (msambare/fer2013)'],
    ['Original Size','35,887 images across 7 classes'],
    ['Classes Used','angry, happy, sad'],
    ['Training Set','~13,665 images  (85% of train split)'],
    ['Validation Set','~2,410 images  (15% of train split)'],
    ['Test Set','3,979 images'],
    ['Resolution','48 × 48 pixels, single-channel grayscale'],
    ['Normalisation','Pixel values rescaled to [0, 1]'],
    ['Augmentation','Rotation ±10°, width/height shift, horizontal flip, zoom 10%'],
],widths=[Cm(5),Cm(9.5)])

h2('2.2 Model Architecture')
para("The architecture follows the Lab 2 tutorial structure pretty closely: three convolutional blocks, each one combining a Conv2D layer with Batch Normalisation, MaxPooling, and Dropout. After flattening there's a single fully connected layer of 256 units, then the three-class softmax output. Adding BatchNorm between the convolution and pooling in each block turned out to make a real difference — training was considerably more stable compared to an earlier run without it.")
code(
    'Input: (48, 48, 1)\n\n'
    'Conv2D(32,  3×3, relu, padding=same) → BatchNorm → MaxPool(2×2) → Dropout(0.25)\n'
    'Conv2D(64,  3×3, relu, padding=same) → BatchNorm → MaxPool(2×2) → Dropout(0.25)\n'
    'Conv2D(128, 3×3, relu, padding=same) → BatchNorm → MaxPool(2×2) → Dropout(0.25)\n\n'
    'Flatten → Dense(256, relu) → BatchNorm → Dropout(0.5) → Dense(3, softmax)\n\n'
    'Optimizer : Adam (default lr=0.001)\n'
    'Loss      : Categorical Crossentropy\n'
    'Callbacks : EarlyStopping (patience=5, restore_best_weights=True)\n'
    '            ReduceLROnPlateau (factor=0.5, patience=3)'
)

h2('2.3 Key Code')
para('Data pipeline — loading and augmentation from emotion_image_cnn.py:')
code(
    'train_datagen = ImageDataGenerator(\n'
    '    rescale=1./255, rotation_range=10,\n'
    '    width_shift_range=0.1, height_shift_range=0.1,\n'
    '    horizontal_flip=True, zoom_range=0.1, validation_split=0.15\n'
    ')\n'
    'train_gen = train_datagen.flow_from_directory(\n'
    '    TRAIN_DIR, target_size=(48, 48), color_mode="grayscale",\n'
    '    classes=["angry", "happy", "sad"], class_mode="categorical",\n'
    '    batch_size=64, subset="training"\n'
    ')'
)
para('Model construction:')
code(
    'model = Sequential([\n'
    '    Conv2D(32,  (3,3), activation="relu", padding="same", input_shape=(48,48,1)),\n'
    '    BatchNormalization(), MaxPooling2D(2,2), Dropout(0.25),\n'
    '    Conv2D(64,  (3,3), activation="relu", padding="same"),\n'
    '    BatchNormalization(), MaxPooling2D(2,2), Dropout(0.25),\n'
    '    Conv2D(128, (3,3), activation="relu", padding="same"),\n'
    '    BatchNormalization(), MaxPooling2D(2,2), Dropout(0.25),\n'
    '    Flatten(), Dense(256, activation="relu"),\n'
    '    BatchNormalization(), Dropout(0.5), Dense(3, activation="softmax")\n'
    '])'
)

h2('2.4 Training Results')
para("Training ran through all 25 epochs without early stopping ever triggering, though around epoch 20 the learning rate did get knocked down by ReduceLROnPlateau when validation loss stopped moving. Final validation accuracy settled somewhere in the 71–72% range. The test set came out a bit higher at 77.71%, which was honestly slightly unexpected at first. It seems to be partly because the validation split came from the training folder distribution, while the test folder has a somewhat different class balance.")
img(os.path.join(IMG_IMAGE,'image_accuracy_loss.png'),w=Inches(5.5),
    c='Figure 2.1: Accuracy (left) and loss (right) curves over training epochs — Image CNN on FER2013.')

h2('2.5 Classification Report and Confusion Matrix')
para('The full classification report on the 3,979-image test set is shown below:')
code(
    '              precision    recall  f1-score   support\n\n'
    '       angry       0.73      0.59      0.65       958\n'
    '       happy       0.84      0.93      0.88      1774\n'
    '         sad       0.72      0.70      0.71      1247\n\n'
    '    accuracy                           0.78      3979\n'
    '   macro avg       0.76      0.74      0.75      3979\n'
    'weighted avg       0.77      0.78      0.77      3979\n\n'
    'Test Accuracy : 77.71%\n'
    'Macro F1      : 0.7462'
)
img(os.path.join(IMG_IMAGE,'image_confusion_matrix.png'),w=Inches(4.5),
    c='Figure 2.2: Confusion matrix on the test set. Happy dominates correct predictions; angry samples frequently land on sad.')

h2('2.6 Sample Prediction')
img(os.path.join(IMG_IMAGE,'image_sample_prediction.png'),w=Inches(3.8),
    c='Figure 2.3: A single test image with its true label, predicted label, and per-class probabilities.')

h2('2.7 Discussion')
para("At 77.71% on a three-class problem the model is clearly picking up on real features — the random baseline would be around 33%, so there's a meaningful gap there. Happy performed best by a noticeable margin, hitting 88% F1. That makes sense when you think about it: a wide open-mouthed smile with raised cheeks looks pretty distinct from the other two expressions, even at low resolution. The angry class was the hardest. Only 59% of angry samples got recalled correctly, with most of the mistakes landing on sad instead.")
para("Looking at the FER2013 images, that confusion is understandable. Both angry and sad expressions involve furrowed brows and downturned mouth corners — the main difference is subtle muscle tension that's genuinely hard to catch at 48×48 pixels. It's probably the resolution more than anything else causing that confusion.")
para("That said, the model doesn't show severe overfitting. Training accuracy was around 74% at the point where validation loss levelled off, so there's no dramatic split between the two curves — the dropout and data augmentation seem to be doing their job. If this were taken further, swapping in a pre-trained backbone like VGGFace or a lightweight MobileNet would probably push accuracy well above 85% even at this resolution.")

# ════════════════════════════════════════════════════════════
#  3. LAB B — AUDIO
# ════════════════════════════════════════════════════════════
doc.add_page_break()
h1('3. Lab B: Emotion Audio Classification (CNN + Chroma-CQT)')

h2('3.1 Dataset Overview')
para("TESS (Toronto Emotional Speech Set) consists of recordings from two actresses — one older adult, one younger — each reading 200 target words in seven different emotional tones, giving 2,800 WAV files in total. For this experiment three emotions were pulled out: angry, happy, and sad. Both speakers have recordings for all three, so it works out to exactly 400 files per class — a perfectly balanced dataset of 1,200 clips. The audio was resampled to 22,050 Hz during feature extraction.")

tbl(['Property','Value'],[
    ['Source','Kaggle  (ejlok1/toronto-emotional-speech-set-tess)'],
    ['Speakers','OAF (older adult female) + YAF (young adult female)'],
    ['Classes Used','angry, happy, sad'],
    ['Files per Class','400  (200 per speaker)'],
    ['Total Files','1,200 WAV clips'],
    ['Sample Rate','44.1 kHz original  →  resampled to 22,050 Hz'],
    ['Feature','Chroma-CQT  (12 bins × 174 frames)'],
    ['Train / Val / Test','840 / 180 / 180  (70 / 15 / 15 %)'],
],widths=[Cm(5),Cm(9.5)])

h2('3.2 Chroma-CQT Feature Extraction')
para("Chroma-CQT maps a signal's energy onto twelve pitch classes (C, C#, D, all the way to B) at each point in time, producing a 12×N matrix where N is the number of time frames. It's compact and has been shown to work well when pitch and harmonic content are the main things that separate the classes — which is exactly why the Lab 3 tutorial used it for musical instrument classification.")
para("Each clip was processed with a hop length of 512 samples. For a typical two-second utterance at 22,050 Hz that gives roughly 174 frames. Clips shorter or longer than that were zero-padded or truncated to land on exactly 174 frames. After that, the resulting 12×174 matrix was normalised per chroma bin — subtracting the bin mean and dividing by the bin standard deviation — before getting passed into the CNN as a single-channel image.")
code(
    'SR=22050, HOP_LENGTH=512, N_CHROMA=12, MAX_FRAMES=174\n\n'
    'y, _ = librosa.load(filepath, sr=SR)\n'
    'chroma = librosa.feature.chroma_cqt(y=y, sr=SR,\n'
    '                                     hop_length=HOP_LENGTH, n_chroma=N_CHROMA)\n'
    '# Per-bin z-score normalisation\n'
    'mu  = chroma.mean(axis=1, keepdims=True)\n'
    'sig = chroma.std(axis=1,  keepdims=True)\n'
    'chroma = (chroma - mu) / (sig + 1e-9)\n'
    '# Pad or truncate to fixed 174 frames, add channel dim → (12, 174, 1)'
)

h2('3.3 Model Architecture')
para("The CNN architecture mirrors the image model, adapted only in the pooling stride of the third block (1×2 instead of 2×2) to account for the narrow height of the 12×174 input:")
code(
    'Input: (12, 174, 1)\n\n'
    'Conv2D(32,  3×3, relu, same) → BatchNorm → MaxPool(2×2) → Dropout(0.25)\n'
    'Conv2D(64,  3×3, relu, same) → BatchNorm → MaxPool(2×2) → Dropout(0.25)\n'
    'Conv2D(128, 3×3, relu, same) → BatchNorm → MaxPool(1×2) → Dropout(0.25)\n\n'
    'Flatten → Dense(256, relu) → BatchNorm → Dropout(0.5) → Dense(3, softmax)\n\n'
    'Optimizer : Adam  |  Loss : Categorical Crossentropy\n'
    'Callbacks : EarlyStopping (patience=7)  +  ReduceLROnPlateau'
)

h2('3.4 Training Results')
para("Training converged in roughly 20 epochs before EarlyStopping called it. Validation accuracy kind of plateaued in the mid-60s, and when run on the test set the final number came out at 68.89%. Not bad, though the per-class story is a lot messier than that single number suggests.")
img(os.path.join(IMG_AUDIO,'audio_accuracy_loss.png'),w=Inches(5.5),
    c='Figure 3.1: Accuracy and loss curves during training — Audio CNN on TESS Chroma-CQT features.')

h2('3.5 Classification Report and Confusion Matrix')
para('Results on the held-out test set of 180 clips (60 per class):')
code(
    '              precision    recall  f1-score   support\n\n'
    '       angry       1.00      0.12      0.21        60\n'
    '       happy       1.00      0.95      0.97        60\n'
    '         sad       0.52      1.00      0.68        60\n\n'
    '    accuracy                           0.69       180\n'
    '   macro avg       0.84      0.69      0.62       180\n'
    'weighted avg       0.84      0.69      0.62       180\n\n'
    'Test Accuracy : 68.89%\n'
    'Macro F1      : 0.6217'
)
img(os.path.join(IMG_AUDIO,'audio_confusion_matrix.png'),w=Inches(4.5),
    c='Figure 3.2: Confusion matrix — Audio CNN. Happy classified near-perfectly; angry almost entirely predicted as sad.')

h2('3.6 Sample Visualisation')
img(os.path.join(IMG_AUDIO,'audio_sample_prediction.png'),w=Inches(5.5),
    c='Figure 3.3: Chroma-CQT feature map of a test sample with true and predicted emotion labels.')

h2('3.7 Discussion')
para("68.89% overall sounds reasonable, but the per-class breakdown tells a very different story. Happy is almost perfectly classified — 95% recall and perfect precision. Sad is never missed at all (100% recall). The problem is that roughly half the angry samples are getting absorbed into the sad predictions, which pulls sad's precision down to 0.52. Angry itself only hits 12% recall, so the model is essentially failing to recognise angry speech and predicting sad almost every time instead.")
para("This is a known limitation of Chroma-CQT for speech emotion — it captures which pitch classes are active and how they shift over time, but it doesn't capture loudness, speaking rate, or voice quality. Those are probably the cues that actually separate angry from sad speech in practice. An MFCC-based feature, or something that combines energy and delta coefficients alongside the chroma representation, would almost certainly recover a big chunk of that lost recall for angry.")
para("One thing still worth pointing out: the precision values for angry and happy are both 1.00. When the model does predict those labels, it's never wrong — it's just that it predicts them far too rarely. That tells you the problem is coverage, not correctness.")

# ════════════════════════════════════════════════════════════
#  4. LAB C — TEXT BERT
# ════════════════════════════════════════════════════════════
doc.add_page_break()
h1('4. Lab C: Emotion Text Classification (BERT)')

h2('4.1 Dataset Overview')
para("The Tweet Emotions dataset has 40,000 tweets annotated across thirteen different emotion labels. Three were selected here: happiness, sadness, and anger, which map reasonably well to the happy/sad/angry classes used in the other two systems.")
para("A practical problem showed up pretty quickly: anger has only 110 labelled tweets in the whole dataset. Happiness has 5,209 and sadness has 5,165. To avoid training on a wildly skewed set, both of those were downsampled to 110 as well, giving a balanced total of 330 tweets — 231 for training, 50 for validation, and 50 for testing. That's a very small dataset for any deep learning model, honestly, and that fact ends up dominating the results more than anything else.")

tbl(['Property','Value'],[
    ['Source','Kaggle  (pashupatigupta/emotion-detection-from-text)'],
    ['Full Dataset','40,000 tweets across 13 emotion labels'],
    ['Classes Used','happiness, sadness, anger'],
    ['Before Balance','happiness: 5,209  |  sadness: 5,165  |  anger: 110'],
    ['After Balance','110 per class  =  330 tweets total'],
    ['Train / Val / Test','231 / 50 / 50  (70 / 15 / 15 %)'],
    ['Backbone','bert-base-uncased  (110M parameters)'],
    ['Max Token Length','128 tokens  (WordPiece tokenisation)'],
],widths=[Cm(5),Cm(9.5)])

h2('4.2 Model Architecture')
para("Following the Lab 5 NLP tutorial, the approach takes the pre-trained BERT base model and attaches a linear classification head on top of the pooled [CLS] output. During fine-tuning everything is trainable — both the transformer blocks and the new classifier head. The model setup is standard HuggingFace:")
code(
    'BertForSequenceClassification\n'
    '  ├─ BertModel\n'
    '  │    ├─ BertEmbeddings      (token + position + segment)\n'
    '  │    ├─ BertEncoder         (12 layers, 768 hidden, 12 heads)\n'
    '  │    └─ BertPooler          (linear + tanh on [CLS] token output)\n'
    '  ├─ Dropout(p=0.1)\n'
    '  └─ Linear(in=768, out=3)    ← newly initialised for this task\n\n'
    'Optimiser  : AdamW  (lr=2e-5, weight_decay=0.01)\n'
    'Scheduler  : Linear warmup over first 10% of steps\n'
    'Epochs     : 5  |  Batch size : 32  |  Grad clip : 1.0'
)

h2('4.3 Key Code')
para('Custom PyTorch Dataset for tokenisation (emotion_text_bert.py):')
code(
    'class EmotionDataset(Dataset):\n'
    '    def __getitem__(self, idx):\n'
    '        enc = self.tokenizer(\n'
    '            str(self.texts[idx]),\n'
    '            max_length=128, padding="max_length",\n'
    '            truncation=True, return_tensors="pt"\n'
    '        )\n'
    '        return {\n'
    '            "input_ids":      enc["input_ids"].squeeze(0),\n'
    '            "attention_mask": enc["attention_mask"].squeeze(0),\n'
    '            "label":          torch.tensor(self.labels[idx])\n'
    '        }'
)
para('Training loop per batch:')
code(
    'outputs = model(input_ids=ids, attention_mask=mask, labels=labels)\n'
    'loss = outputs.loss\n'
    'loss.backward()\n'
    'torch.nn.utils.clip_grad_norm_(model.parameters(), 1.0)\n'
    'optimizer.step(); scheduler.step(); optimizer.zero_grad()'
)

h2('4.4 Training Results')
para('Epoch-level training and validation metrics over five epochs:')
code(
    'Epoch   Train Loss   Train Acc   Val Loss   Val Acc\n'
    '─────   ──────────   ─────────   ────────   ───────\n'
    '  1       1.1305      36.36%     1.0905     46.94%\n'
    '  2       1.0775      43.72%     1.1214     42.86%\n'
    '  3       1.0348      54.11%     1.0843     42.86%\n'
    '  4       1.0130      56.71%     1.0768     46.94%\n'
    '  5       0.9820      63.64%     1.0767     44.90%'
)
img(os.path.join(IMG_NLP,'text_accuracy_loss.png'),w=Inches(5.5),
    c='Figure 4.1: Training and validation accuracy/loss across five epochs — BERT on Tweet Emotions.')

h2('4.5 Classification Report and Confusion Matrix')
para('Test set results on 50 tweets:')
code(
    '              precision    recall  f1-score   support\n\n'
    '   happiness       0.61      0.69      0.65        16\n'
    '     sadness       0.53      0.53      0.53        17\n'
    '       anger       0.40      0.35      0.38        17\n\n'
    '    accuracy                           0.52        50\n'
    '   macro avg       0.51      0.52      0.52        50\n'
    'weighted avg       0.51      0.52      0.51        50\n\n'
    'Test Accuracy : 52.00%\n'
    'Macro F1      : 0.5172'
)
img(os.path.join(IMG_NLP,'text_confusion_matrix.png'),w=Inches(4.5),
    c='Figure 4.2: Confusion matrix — BERT test set. Happiness identified reasonably well; anger most confused.')

h2('4.6 Sample Prediction')
para('One example from the test set:')
code(
    'Input  : "@olafsearson Lol - I could try!  Seriously tho, dont do all of it! xx"\n'
    'True   : anger\n'
    'Pred   : happiness\n\n'
    'happiness : 0.371\n'
    'sadness   : 0.284\n'
    'anger     : 0.346'
)
para("This is a sarcastic tweet, and the model nearly gets it right. Anger scores 0.346 against happiness at 0.371 — they're almost tied. With more training data that covers a wider range of sarcastic patterns, the model would probably resolve that ambiguity correctly.")

h2('4.7 Discussion')
para("52% on a three-class task is above the 33% chance level, so something is being learned here, but it's clearly not enough for practical use. The reason isn't complicated: 231 training samples is nowhere near sufficient to meaningfully fine-tune a 110M parameter model. BERT's pre-training gives it a strong foundation in English semantics, but picking up on the emotional nuances in short, informal tweets requires task-specific patterns the model just hasn't seen enough of.")
para("The anger class is the real limiting factor. Even before balancing, there were only 110 anger tweets in the entire 40,000-tweet dataset — that's a thin slice. The validation curves also jump around quite a bit between epochs, fluctuating between roughly 42% and 47%, which is a direct consequence of having only 50 validation samples. One misclassification shifts accuracy by two percentage points, so it's hard to read too much into that curve.")
para("That said, the approach itself seems sound. With something like GoEmotions or the SemEval 2018 emotion task — which give thousands of examples per class — the same fine-tuning setup would typically reach around 85–92% accuracy. The architecture isn't the bottleneck here. The data is.")

# ════════════════════════════════════════════════════════════
#  5. CROSS-MODALITY COMPARISON
# ════════════════════════════════════════════════════════════
doc.add_page_break()
h1('5. Cross-Modality Comparison')

h2('5.1 Summary Table')
tbl(['Modality','Dataset','Architecture','Train Samples','Test Acc.','Macro F1'],[
    ['Image','FER2013',       'CNN  (3 conv blocks)', '~13,665','77.71%','0.75'],
    ['Audio','TESS',          'CNN + Chroma-CQT',      '840',   '68.89%','0.62'],
    ['Text', 'Tweet Emotions','BERT fine-tuned',        '231',  '52.00%','0.52'],
],widths=[Cm(2.2),Cm(3.5),Cm(3.5),Cm(2.5),Cm(2.3),Cm(2.0)])

h2('5.2 What the Numbers Show')
para("The clearest pattern across the three systems is that accuracy tracks almost perfectly with training set size. The image model, with roughly 13,600 training examples, scores nearly 78%. Audio, with 840 examples, sits at 69%. Text, with just 231, reaches only 52%. That's probably not because images are inherently easier than text for emotion classification — it's more that data availability just happened to be very different across the three modalities in this particular experiment.")
para("The one result that breaks that pattern is the audio happy class, which hits a 0.97 F1 score even though the overall model only scores 0.62 macro F1. Happy speech has a characteristic upward pitch contour and rhythmic variety that shows up clearly in Chroma-CQT. Angry and sad, on the other hand, don't have pitch signatures that the feature can separate well — which explains why those classes drag the overall number down.")

h2('5.3 Per-Class Results Across Modalities')
tbl(['Emotion','Image F1','Audio F1','Text F1','Pattern'],[
    ['Happy / Happiness','0.88','0.97','0.65','Easiest in all three — most visually and acoustically distinct'],
    ['Sad / Sadness',    '0.71','0.68','0.53','Middle range — consistent across modalities'],
    ['Angry / Anger',    '0.65','0.21','0.38','Hardest in audio and text; image does OK'],
],widths=[Cm(3.3),Cm(1.8),Cm(1.8),Cm(1.8),Cm(5.7)])

para("Angry is the weak point across all three systems, though for different reasons in each one. In audio, the Chroma-CQT feature just can't pull it apart from sad. In text, there simply weren't enough samples to learn the class properly. Only in images does angry perform reasonably — because the facial muscle tension that comes with it, the furrowed brows and tightened lips, is visible enough at 48×48 for the CNN to partially pick up on. Still the hardest of the three image classes, but at least learnable.")

h2('5.4 Architecture Observations')
para("All three CNN models share the same three-block structure from the Lab 2/3 tutorials. That was intentional — keeping the architecture consistent means any differences in results come from the feature representation and the dataset, not from swapping out model designs. BERT is a much larger, more powerful model by parameter count, though model capacity clearly isn't the limiting factor when you're working with this little data.")
para("If data volumes were equalised across all three modalities, I'd expect BERT and the image CNN to compete closely at the top. The audio CNN would probably trail somewhat unless the feature extraction gets upgraded beyond Chroma-CQT.")

# ════════════════════════════════════════════════════════════
#  6. CONCLUSION
# ════════════════════════════════════════════════════════════
doc.add_page_break()
h1('6. Conclusion')

para("This project built and evaluated three separate emotion detection systems — one for images, one for audio, one for text — all targeting the same three emotion classes. The results came out in a clear gradient: image classification at the top at 77.71%, audio in the middle at 68.89%, and text at the bottom at 52.00%. As the cross-modality analysis makes clear, that gradient pretty closely reflects training dataset size rather than any fundamental difference in how hard each modality is.")

para("The image CNN on FER2013 gave the most reliable results overall. The dataset is large and reasonably well-curated, and the Lab 2 architecture worked well for the task after adding BatchNormalization and data augmentation. The main remaining weakness is the angry class, which shares too many low-level visual features with sad to be reliably separated at 48×48 resolution.")

para("The audio CNN on TESS revealed a specific limitation of Chroma-CQT: it captures pitch content but not vocal intensity or quality, and those acoustic markers are exactly what distinguishes angry speech from sad. The happy class result (0.97 F1) shows the feature working well when it's informative, and the angry collapse (0.21 F1) shows where it breaks down. Switching to MFCCs or a combined feature set is probably the most logical first fix.")

para("The BERT text model ran into a data problem more than anything else. 231 training samples isn't enough to fine-tune a 110M parameter model for something as nuanced as informal tweet classification. The approach is sound — loss did go down and accuracy did improve through all five epochs — it just didn't go far enough to reach practically useful performance at this scale. With a larger anger-labelled dataset like GoEmotions, the same setup would likely do well.")

para("Taken together, all three experiments point toward the same takeaway: labelled training data tends to matter more than architecture, at least in the low-data regime. A multimodal system combining face, voice, and language would be a natural next step and would likely outperform any single-modality model, especially on the ambiguous cases that trip up individual systems.")

para('Project file structure:')
code(
    'emotion_detection_submission/\n'
    '  lab_image/   emotion_image_cnn.py  |  image_accuracy_loss.png\n'
    '               image_confusion_matrix.png  |  image_sample_prediction.png\n'
    '  lab_audio/   emotion_audio_cnn.py  |  audio_accuracy_loss.png\n'
    '               audio_confusion_matrix.png  |  audio_sample_prediction.png\n'
    '  lab_nlp/     emotion_text_bert.py  |  text_accuracy_loss.png\n'
    '               text_confusion_matrix.png\n'
    '  Advanced_HCI_Phase2_Report.docx'
)

out = os.path.join(BASE, 'Advanced_HCI_Phase2_Report_FINAL.docx')
doc.save(out)
print('Saved:', out)
